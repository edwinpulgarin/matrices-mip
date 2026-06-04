from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "metodologia_avances_siguientes_pasos.md"
OUT_PATH = ROOT / "docs" / "Metodologia_Avances_Siguientes_Pasos_MIP.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=10.5, bold=False, color=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def add_inline(paragraph, text, base_bold=False, size=10.5):
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, bold=base_bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=9, color="1F4E79", font="Consolas")
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=base_bold)


def style_paragraph(paragraph, space_after=6):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.08


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7FA")
    set_cell_border(cell, "D9E2EC")
    set_cell_margins(cell, 110, 130, 110, 130)
    paragraph = cell.paragraphs[0]
    for line in lines:
        run = paragraph.add_run(line + "\n")
        set_run_font(run, size=9, color="404040", font="Consolas")
    doc.add_paragraph()


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(cell.replace(":", "").replace("-", "").strip() == "" for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(rows):
        set_row_cant_split(table.rows[row_idx])
        if row_idx == 0:
            set_repeat_table_header(table.rows[row_idx])
        for col_idx in range(n_cols):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            value = row[col_idx] if col_idx < len(row) else ""
            add_inline(paragraph, value, base_bold=(row_idx == 0), size=9.3)
            if row_idx == 0:
                set_cell_shading(cell, "DDEBF7")
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string("1F4E79")
                    run.bold = True
            elif row_idx % 2 == 0:
                set_cell_shading(cell, "F7FBFD")
    doc.add_paragraph()


def add_numbered_item(doc, number, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.22)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{number}. ")
    set_run_font(run, bold=True)
    add_inline(paragraph, text)


def build_document():
    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Heading 1", 17, "1F4E79"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 12, "2F6F73"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    run = title.add_run("Metodología, avances y siguientes pasos\n")
    set_run_font(run, size=25, bold=True, color="1F4E79", font="Arial")
    run = title.add_run("Repositorio de matrices insumo-producto")
    set_run_font(run, size=16, color="2F6F73", font="Arial")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(28)
    for line in [
        "Argentina, Brasil, México y Uruguay",
        "Actualización: 29 de mayo de 2026",
        "Basado en la metodología de MIP extendida ambientalmente y huella de carbono",
    ]:
        run = meta.add_run(line + "\n")
        set_run_font(run, size=11, color="505050")

    doc.add_page_break()

    lines = text.splitlines()
    idx = 0
    in_code = False
    code_lines = []
    ordered_restart = True
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if not stripped:
            ordered_restart = True
            idx += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            add_table(doc, parse_table(table_lines))
            ordered_restart = True
            continue
        if stripped.startswith("# "):
            idx += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:]
            if heading.startswith(("8. ", "12. ")):
                doc.add_page_break()
            paragraph = doc.add_heading(heading, level=1)
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(4)
            ordered_restart = True
            idx += 1
            continue
        if stripped.startswith("### "):
            paragraph = doc.add_heading(stripped[4:], level=2)
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
            ordered_restart = True
            idx += 1
            continue
        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            style_paragraph(paragraph, 3)
            idx += 1
            continue
        match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if match:
            add_numbered_item(doc, match.group(1), match.group(2))
            idx += 1
            continue
        paragraph = doc.add_paragraph()
        add_inline(paragraph, stripped)
        style_paragraph(paragraph)
        idx += 1

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Repositorio MIP | Metodología, avances y siguientes pasos")
        set_run_font(run, size=8, color="808080")

    OUT_PATH.parent.mkdir(exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
